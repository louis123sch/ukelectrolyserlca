from __future__ import annotations

import io
import mimetypes
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Iterator

import fitz  # PyMuPDF
from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook


@dataclass(frozen=True)
class VisualAsset:
    """One image/page supplied to the vision evidence-extraction stage."""

    document: str
    asset_id: str
    mime_type: str
    data: bytes
    context: str | None = None
    source_type: str = "figure"


_RASTER_MIME_TYPES = {"image/png", "image/jpeg", "image/webp", "image/gif"}
_VECTOR_EXTENSIONS = {".emf", ".wmf", ".svg"}
_VISUAL_RELEVANCE_TERMS = (
    "inventory",
    "lci",
    "input",
    "output",
    "material",
    "component",
    "ecoinvent",
    "system boundary",
    "process",
    "values",
    "table",
    "flow",
    "mass",
    "energy",
)


def extract_pdf_text(pdf_bytes: bytes, max_pages: int | None = None) -> str:
    """Extract page-tagged text from a text-readable PDF locally."""
    if not pdf_bytes:
        raise ValueError("PDF is empty")

    doc = fitz.open(stream=io.BytesIO(pdf_bytes), filetype="pdf")
    chunks: list[str] = []
    n_pages = len(doc) if max_pages is None else min(len(doc), max_pages)

    for idx in range(n_pages):
        page = doc[idx]
        text = page.get_text("text").strip()
        chunks.append(f"[PAGE {idx + 1}]\n{text}")

    combined = "\n\n".join(chunks).strip()
    if not combined:
        raise ValueError(
            "No machine-readable text was found. This PDF may be scanned/image-only; "
            "multimodal page rendering is required."
        )
    return combined


def extract_pdf_visual_assets(
    pdf_bytes: bytes,
    *,
    filename: str = "document.pdf",
    max_pages: int | None = None,
) -> list[VisualAsset]:
    """Render pages that contain embedded images, or pages with very little readable text.

    This is intentionally selective: native PDF text remains the primary evidence source,
    while visually rich/scanned pages are made available to a vision model.
    """
    if not pdf_bytes:
        return []
    doc = fitz.open(stream=io.BytesIO(pdf_bytes), filetype="pdf")
    n_pages = len(doc) if max_pages is None else min(len(doc), max_pages)
    assets: list[VisualAsset] = []
    for idx in range(n_pages):
        page = doc[idx]
        text = page.get_text("text").strip()
        has_images = bool(page.get_images(full=True))
        looks_scanned = len(text) < 80
        if not (has_images or looks_scanned):
            continue
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
        assets.append(
            VisualAsset(
                document=filename,
                asset_id=f"page-{idx + 1}",
                mime_type="image/png",
                data=pix.tobytes("png"),
                context=text[:1200] or "Page contains little/no machine-readable text.",
                source_type="pdf_page",
            )
        )
    return assets


def _iter_docx_blocks(document: _Document) -> Iterator[Paragraph | Table]:
    """Yield DOCX paragraphs and tables in document order."""
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def extract_docx_text(docx_bytes: bytes) -> str:
    """Extract text from a DOCX while retaining paragraph/table provenance markers."""
    if not docx_bytes:
        raise ValueError("DOCX is empty")

    document = Document(io.BytesIO(docx_bytes))
    chunks: list[str] = []
    paragraph_no = 0
    table_no = 0

    for block in _iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                paragraph_no += 1
                chunks.append(f"[PARAGRAPH {paragraph_no}]\n{text}")
        else:
            table_no += 1
            rows = []
            for row in block.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append("\t".join(cells))
            if rows:
                chunks.append(f"[TABLE {table_no}]\n" + "\n".join(rows))

    combined = "\n\n".join(chunks).strip()
    if not combined:
        raise ValueError("No readable text was found in the DOCX")
    return combined


def extract_xlsx_text(xlsx_bytes: bytes) -> str:
    """Extract sheet text from an Excel workbook while retaining sheet provenance markers."""
    if not xlsx_bytes:
        raise ValueError("Excel workbook is empty")

    try:
        workbook = load_workbook(io.BytesIO(xlsx_bytes), data_only=True, read_only=True)
    except Exception as exc:
        raise ValueError("Could not read the Excel workbook. Use .xlsx or .xlsm.") from exc

    chunks: list[str] = []
    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if value is None else str(value).strip() for value in row]
            if any(cells):
                rows.append("\t".join(cells))
        if rows:
            chunks.append(f"[SHEET: {sheet.title}]\n" + "\n".join(rows))

    combined = "\n\n".join(chunks).strip()
    if not combined:
        raise ValueError("No readable data was found in the Excel workbook")
    return combined


def _find_office_converter() -> str | None:
    candidates = [
        shutil.which("libreoffice"),
        shutil.which("soffice"),
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ]
    return next((str(x) for x in candidates if x and Path(x).exists()), None)


def _convert_vector_to_png(data: bytes, suffix: str) -> bytes | None:
    """Rasterize Office vector artwork (EMF/WMF/SVG) when LibreOffice is available."""
    converter = _find_office_converter()
    if not converter:
        return None
    with tempfile.TemporaryDirectory(prefix="ai_lca_visual_") as tmp:
        src = Path(tmp) / f"source{suffix}"
        src.write_bytes(data)
        result = subprocess.run(
            [converter, "--headless", "--convert-to", "png", "--outdir", tmp, str(src)],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=60,
            check=False,
        )
        png = Path(tmp) / "source.png"
        if result.returncode != 0 or not png.exists():
            return None
        return png.read_bytes()


def _nearby_docx_context(paragraphs: list[Paragraph], index: int) -> str:
    snippets: list[str] = []
    for step in (-2, -1, 1, 2):
        j = index + step
        if 0 <= j < len(paragraphs):
            text = paragraphs[j].text.strip()
            if text:
                snippets.append(text)
    return " | ".join(snippets[:3])[:1200]


def extract_docx_visual_assets(
    docx_bytes: bytes,
    *,
    filename: str = "document.docx",
) -> tuple[list[VisualAsset], list[str]]:
    """Extract embedded DOCX figures with nearby caption/prose context.

    Raster formats are passed through directly. Common Office vector formats such as
    EMF/WMF are rasterized with LibreOffice when available. Unsupported vector media
    are reported explicitly rather than silently ignored.
    """
    if not docx_bytes:
        return [], []
    document = Document(io.BytesIO(docx_bytes))
    paragraphs = list(document.paragraphs)
    seen_rids: set[str] = set()
    assets: list[VisualAsset] = []
    warnings: list[str] = []

    def add_part(rid: str, context: str) -> None:
        if not rid or rid in seen_rids:
            return
        part = document.part.related_parts.get(rid)
        if part is None or not getattr(part, "blob", None):
            return
        seen_rids.add(rid)
        blob = part.blob
        partname = str(getattr(part, "partname", "image"))
        suffix = Path(partname).suffix.lower()
        mime = getattr(part, "content_type", None) or mimetypes.guess_type(partname)[0] or "application/octet-stream"
        if mime in _RASTER_MIME_TYPES:
            out_data, out_mime = blob, mime
        elif suffix in _VECTOR_EXTENSIONS or mime in {"image/x-emf", "image/x-wmf", "image/svg+xml"}:
            converted = _convert_vector_to_png(blob, suffix or ".emf")
            if converted is None:
                warnings.append(
                    f"{filename}: embedded vector figure {partname} could not be rasterized. "
                    "Install LibreOffice to make EMF/WMF/SVG figures available to vision."
                )
                return
            out_data, out_mime = converted, "image/png"
        else:
            warnings.append(f"{filename}: unsupported embedded image format {mime} ({partname}).")
            return
        assets.append(
            VisualAsset(
                document=filename,
                asset_id=f"figure-{len(assets) + 1}",
                mime_type=out_mime,
                data=out_data,
                context=context or None,
                source_type="docx_figure",
            )
        )

    for idx, paragraph in enumerate(paragraphs):
        for blip in paragraph._p.xpath(".//a:blip"):
            rid = blip.get(qn("r:embed"))
            add_part(rid, _nearby_docx_context(paragraphs, idx))

    # Include any image relationships not represented by document.paragraphs (e.g. tables).
    for rid, rel in document.part.rels.items():
        if "image" in str(getattr(rel, "reltype", "")):
            add_part(rid, "Embedded image; no nearby paragraph context was available.")

    return assets, warnings


def rank_visual_assets(assets: Iterable[VisualAsset], *, limit: int = 24) -> list[VisualAsset]:
    """Prioritize images whose nearby text suggests LCI/model evidence."""
    def score(asset: VisualAsset) -> tuple[int, int]:
        context = (asset.context or "").casefold()
        hits = sum(term in context for term in _VISUAL_RELEVANCE_TERMS)
        return hits, len(context)

    ranked = sorted(list(assets), key=score, reverse=True)
    return ranked[: max(0, limit)]


def extract_text_fixture(text_bytes: bytes) -> str:
    """Read a UTF-8 text fixture used by private regression benchmarks."""
    if not text_bytes:
        raise ValueError("Text fixture is empty")
    try:
        text = text_bytes.decode("utf-8").strip()
    except UnicodeDecodeError as exc:
        raise ValueError("Text fixture must be UTF-8 encoded") from exc
    if not text:
        raise ValueError("No readable text was found in the text fixture")
    return text


def extract_document_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch supported source documents to the appropriate local text extractor."""
    suffix = Path(filename or "").suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(file_bytes)
    if suffix == ".docx":
        return extract_docx_text(file_bytes)
    if suffix == ".txt":
        return extract_text_fixture(file_bytes)
    if suffix in {".xlsx", ".xlsm"}:
        return extract_xlsx_text(file_bytes)
    raise ValueError("Unsupported document type. Use PDF, DOCX, XLSX, or UTF-8 TXT.")


def extract_document_visual_assets(file_bytes: bytes, filename: str) -> tuple[list[VisualAsset], list[str]]:
    suffix = Path(filename or "").suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_visual_assets(file_bytes, filename=filename), []
    if suffix == ".docx":
        return extract_docx_visual_assets(file_bytes, filename=filename)
    return [], []


def combine_document_texts(documents: Iterable[tuple[str, bytes]]) -> str:
    """Extract and combine multiple source documents with explicit document markers."""
    chunks: list[str] = []
    for filename, file_bytes in documents:
        filename = (filename or "source").strip() or "source"
        text = extract_document_text(file_bytes, filename)
        chunks.append(f"[DOCUMENT: {filename}]\n{text}")
    combined = "\n\n".join(chunks).strip()
    if not combined:
        raise ValueError("No readable source documents were supplied")
    return combined


def combine_document_evidence(
    documents: Iterable[tuple[str, bytes]],
    *,
    max_visual_assets: int = 24,
) -> tuple[str, list[VisualAsset], list[str]]:
    """Return native text plus selected visual assets and ingestion warnings."""
    items = list(documents)
    text = combine_document_texts(items)
    visual_assets: list[VisualAsset] = []
    warnings: list[str] = []
    for filename, file_bytes in items:
        assets, asset_warnings = extract_document_visual_assets(file_bytes, filename)
        visual_assets.extend(assets)
        warnings.extend(asset_warnings)
    return text, rank_visual_assets(visual_assets, limit=max_visual_assets), warnings
